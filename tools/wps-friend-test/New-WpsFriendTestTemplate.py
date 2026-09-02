import sys
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


def set_font(run, size=11, bold=False):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Microsoft YaHei")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold


def main(output_path):
    document = Document()
    section = document.sections[0]
    section.top_margin = Pt(56)
    section.bottom_margin = Pt(56)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(title.add_run("WPS 集成测试报告模板"), 16, True)
    intro = document.add_paragraph()
    set_font(intro.add_run("本模板仅用于 WPS 加载项验收，所有内容均为虚构测试数据。"), 10)

    for field in ["项目名称", "客户名称", "逆变器型号", "装机容量", "测试日期"]:
        paragraph = document.add_paragraph()
        set_font(paragraph.add_run(f"{field}：XXXXXXXX"), 11)

    note = document.add_paragraph()
    set_font(note.add_run("提示：每个 XXXXXXXX 已由 OOXML 批注定义为对应字段。"), 9)
    document.save(output_path)


if __name__ == "__main__":
    if len(sys.argv) != 2:
        raise SystemExit("usage: New-WpsFriendTestTemplate.py <output.docx>")
    main(sys.argv[1])
